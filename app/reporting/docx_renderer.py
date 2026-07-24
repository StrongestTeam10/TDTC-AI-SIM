"""분석 결과를 수정 가능한 DOCX 정책 사전검토 초안으로 렌더링한다."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .narrative import EvidenceNote, Narrative
from app.schemas.report_models import EvidenceItem, ReportRequest


class DocxReportRenderer:
    def __init__(self) -> None:
        self.font_name = os.getenv("DOCX_FONT_NAME", "맑은 고딕")

    def render(
        self,
        request: ReportRequest,
        narrative: Narrative,
        evidence: list[EvidenceItem],
        charts: dict[str, str],
        output_path: Path,
    ) -> Path:
        document = Document()
        self._configure(document)
        self._cover(document, request, narrative)
        self._toc(document, charts, evidence)
        self._summary(document, narrative)
        self._overview(document, request)
        self._scenario_table(document, request)
        self._result_table(document, request)
        if charts:
            self._charts(document, charts)
        self._flow_and_analysis(
            document,
            narrative,
        )
        self._recommendation(document, narrative)
        self._limitations(document, narrative)
        if evidence:
            self._evidence(
                document,
                narrative.evidence_notes,
            )
        self._appendix(document)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            document.save(output_path)
        except PermissionError as exc:
            raise PermissionError(
                "기존 DOCX 파일을 덮어쓸 수 없습니다. "
                "Word 또는 한컴오피스에서 다음 파일을 닫은 뒤 다시 실행하세요: "
                f"{output_path}"
            ) from exc
        return output_path

    def _configure(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        styles = document.styles
        for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
            style = styles[style_name]
            style.font.name = self.font_name
            style._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)
        styles["Normal"].font.size = Pt(10.5)
        styles["Heading 1"].font.size = Pt(16)
        styles["Heading 2"].font.size = Pt(13)

    def _set_cell_text(self, cell, text: str, bold: bool = False) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = self.font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)
        run.font.size = Pt(9.5)
        cell.vertical_alignment = 1

    @staticmethod
    def _shade(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def _add_heading(self, document: Document, text: str, level: int = 1) -> None:
        paragraph = document.add_heading(text, level=level)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    def _add_bullets(
        self,
        document: Document,
        items: Iterable[str] | str | None,
    ) -> None:
        """문자열 또는 문자열 목록을 불릿 문단으로 추가한다."""
        if items is None:
            return
        normalized_items = [items] if isinstance(items, str) else items

        for item in normalized_items:
            text = str(item).strip()

            if not text:
                continue
            paragraph = document.add_paragraph(
                style="List Bullet"
            )
            run = paragraph.add_run(text)

            # 불릿 본문에도 한글 글꼴 적용
            run.font.name = self.font_name
            run._element.rPr.rFonts.set(
                qn("w:eastAsia"),
                self.font_name,
            )

    def _cover(
        self,
        document: Document,
        request: ReportRequest,
        narrative: Narrative,
    ) -> None:
        for _ in range(4):
            document.add_paragraph()
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(narrative.report_title)
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = self.font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("디지털 트윈 시뮬레이션 기반 정책변경 사전검토 초안").font.size = Pt(13)
        document.add_paragraph()
        table = document.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        values = [
            ("보고서 ID", request.report_id),
            ("정책변경 그룹", str(request.change_id)),
            ("대상 시장", request.market.market_name),
            ("작성 기준일", request.context.analysis_date.isoformat()),
            ("검토 질문", request.decision_question),
        ]
        for row, (label, value) in zip(table.rows, values):
            self._set_cell_text(row.cells[0], label, True)
            self._shade(row.cells[0], "E7EEF8")
            self._set_cell_text(row.cells[1], value)
        document.add_paragraph()
        note = document.add_paragraph(request.disclaimer)
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in note.runs:
            run.font.size = Pt(9)
        document.add_page_break()

    def _toc(self, document: Document, charts: dict[str, str], evidence: list[EvidenceItem]) -> None:
        self._add_heading(document, "목차", 1)
        sections = [
            "1. 검토 결과 요약",
            "2. 검토 개요",
            "3. 시나리오 구성",
            "4. 시나리오별 예측 결과",
        ]
        if charts:
            sections.append("5. 예측 결과 시각화")
        sections.extend(
            [
                "6. 유동 방향 및 추가 분석",
                "7. 종합 검토 의견",
                "8. 분석 한계",
            ]
        )
        if evidence:
            sections.append("9. 검색 근거자료")
        sections.append("붙임. 핵심지표 정의")
        for section in sections:
            document.add_paragraph(section)
        document.add_page_break()

    def _summary(self, document: Document, narrative: Narrative) -> None:
        self._add_heading(document, "1. 검토 결과 요약", 1)
        document.add_paragraph(narrative.executive_summary)

    def _overview(self, document: Document, request: ReportRequest) -> None:
        self._add_heading(document, "2. 검토 개요", 1)
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        items = [
            ("대상 시장", request.market.market_name),
            (
                "시뮬레이션 기간",
                f"{request.context.simulation_start} ~ "
                f"{request.context.simulation_end}",
            ),
            ("예상 방문객", f"{request.context.expected_visitors:,}명"),
            ("모델 버전", request.context.model_version or "미입력"),
        ]
        if (
            request.market.latitude is not None
            and request.market.longitude is not None
        ):
            items.insert(
                1,
                (
                    "시장 위치",
                    f"{request.market.latitude}, "
                    f"{request.market.longitude}",
                ),
            )
        for label, value in items:
            cells = table.add_row().cells
            self._set_cell_text(cells[0], label, True)
            self._shade(cells[0], "E7EEF8")
            self._set_cell_text(cells[1], str(value))
        if request.context.assumptions:
            self._add_heading(document, "분석 가정", 2)
            self._add_bullets(document, request.context.assumptions)

    def _scenario_table(self, document: Document, request: ReportRequest) -> None:
        self._add_heading(document, "3. 시나리오 구성", 1)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["구분", "시나리오명", "설명", "주요 변경사항"]
        for cell, header in zip(table.rows[0].cells, headers):
            self._set_cell_text(cell, header, True)
            self._shade(cell, "D9E5F6")
        for index, alternative in enumerate([request.baseline, *request.alternatives]):
            cells = table.add_row().cells
            interventions = "; ".join(item.description for item in alternative.interventions) or "변경 없음"
            values = [
                "기준안" if index == 0 else f"대안 {index}",
                alternative.alternative_name,
                alternative.description or "-",
                interventions,
            ]
            for cell, value in zip(cells, values):
                self._set_cell_text(cell, str(value))

    @staticmethod
    def _fmt(value, digits: int = 1) -> str:
        if value is None:
            return "미산출"
        if isinstance(value, float):
            return f"{value:.{digits}f}"
        return str(value)

    def _result_table(
        self,
        document: Document,
        request: ReportRequest,
    ) -> None:
        """전달받은 시뮬레이션 결과를 시나리오별로 출력한다."""

        self._add_heading(
            document,
            "4. 시나리오별 예측 결과",
            1,
        )

        table = document.add_table(
            rows=1,
            cols=5,
        )
        table.style = "Table Grid"

        headers = [
            "시나리오",
            "최대 밀집도",
            "평균 밀집도",
            "위험점수",
            "평균 체류시간",
        ]

        for cell, header in zip(
            table.rows[0].cells,
            headers,
        ):
            self._set_cell_text(
                cell,
                header,
                True,
            )
            self._shade(
                cell,
                "D9E5F6",
            )

        scenarios = [
            request.baseline,
            *request.alternatives,
        ]

        for alternative in scenarios:
            metrics = alternative.metrics

            values = [
                alternative.alternative_name,
                self._fmt(metrics.max_density_p_m2),
                self._fmt(metrics.avg_density_p_m2),
                self._fmt(metrics.risk_score, 0),
                self._fmt(metrics.avg_dwell_time_min),
            ]

            cells = table.add_row().cells

            for cell, value in zip(cells, values):
                self._set_cell_text(
                    cell,
                    str(value),
                )

    def _charts(self, document: Document, charts: dict[str, str]) -> None:
        self._add_heading(document, "5. 예측 결과 시각화", 1)
        for key in [
            "density",
            "density_timeseries",
            "risk",
        ]:
            path = charts.get(key)
            if path and Path(path).exists():
                document.add_picture(path, width=Cm(16.0))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _flow_and_analysis(
        self,
        document: Document,
        narrative: Narrative,
    ) -> None:
        """LLM이 원본 유동 데이터를 해석한 시나리오별 문장을 출력한다."""

        self._add_heading(document, "6. 유동 방향 및 추가 분석", 1)
        for item in narrative.flow_analysis:
            self._add_heading(
                document,
                item.alternative_name,
                2,
            )
            document.add_paragraph(
                item.analysis
            )

    def _recommendation(
        self,
        document: Document,
        narrative: Narrative,
    ) -> None:
        self._add_heading(document, "7. 종합 검토 의견", 1)
        document.add_paragraph(narrative.current_issue)
        self._add_heading(document, "검토 근거", 2)
        document.add_paragraph(narrative.recommendation_rationale)
        self._add_heading(document, "후속 조치", 2)
        self._add_bullets(document, narrative.implementation_plan)

    def _limitations(self, document: Document, narrative: Narrative) -> None:
        self._add_heading(document, "8. 분석 한계", 1)
        self._add_bullets(document, narrative.limitations)

    def _evidence(
        self,
        document: Document,
        evidence_notes: list[EvidenceNote],
    ) -> None:
        self._add_heading(document, "9. 검색 근거자료", 1)
        for index, item in enumerate(
            evidence_notes,
            start=1,
        ):
            page = f", {item.page}쪽" if item.page else ""
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"{index}. {item.title}{page}")
            run.bold = True
            document.add_paragraph(item.summary)

    def _appendix(self, document: Document) -> None:
        document.add_section(WD_SECTION.NEW_PAGE)
        self._add_heading(document, "붙임. 핵심지표 정의", 1)
        rows = [
            ("최대 밀집도", "시뮬레이션 기간 중 관측된 최대 단위면적당 인원"),
            ("평균 밀집도", "시뮬레이션 기간의 평균 단위면적당 인원"),
            (
                "예측 위험점수",
                "멀티모달 위험도 분석 시스템에서 산출되어 "
                "보고서 엔진에 전달된 위험도 값",
            ),
            ("평균 체류시간", "방문객이 시장 안에 머문 평균 시간"),
        ]
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        self._set_cell_text(table.rows[0].cells[0], "지표", True)
        self._set_cell_text(table.rows[0].cells[1], "정의", True)
        self._shade(table.rows[0].cells[0], "D9E5F6")
        self._shade(table.rows[0].cells[1], "D9E5F6")
        for label, definition in rows:
            cells = table.add_row().cells
            self._set_cell_text(cells[0], label)
            self._set_cell_text(cells[1], definition)
